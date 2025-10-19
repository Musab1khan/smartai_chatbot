"""
Export Manager - Export reports to Excel, Word, PDF formats
"""

import frappe
import pandas as pd
import logging
from typing import Dict, List, Any
from datetime import datetime
import os

logger = logging.getLogger(__name__)


class ExportManager:
    """Handle report exports in multiple formats"""
    
    @frappe.whitelist()
    def export_to_excel(self, data: List[Dict], filename: str, sheet_name: str = "Report") -> Dict:
        """Export data to Excel format"""
        try:
            df = pd.DataFrame(data)
            
            # Generate file path
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_filename = "".join(c for c in filename if c.isalnum() or c in (' ', '_', '-'))
            file_path = f"/tmp/{safe_filename}_{timestamp}.xlsx"
            
            # Create Excel file
            with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name=sheet_name, index=False)
                
                # Add formatting
                workbook = writer.book
                worksheet = writer.sheets[sheet_name]
                
                # Auto-adjust column widths
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(cell.value)
                        except:
                            pass
                    worksheet.column_dimensions[column_letter].width = max_length + 2
            
            return {
                "success": True,
                "file_path": file_path,
                "filename": os.path.basename(file_path),
                "format": "xlsx",
                "rows": len(df),
                "columns": len(df.columns)
            }
        
        except Exception as e:
            logger.error(f"Error exporting to Excel: {str(e)}")
            return {"success": False, "error": str(e)}
    
    @frappe.whitelist()
    def export_to_word(self, title: str, data: List[Dict], filename: str) -> Dict:
        """Export data to Word document"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            # Create document
            doc = Document()
            
            # Add title
            title_paragraph = doc.add_heading(title, level=1)
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add metadata
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            meta_para.runs[0].font.italic = True
            
            doc.add_paragraph()  # Empty line
            
            # Add table
            if data:
                table = doc.add_table(rows=1, cols=len(data[0]))
                table.style = 'Light Grid Accent 1'
                
                # Add headers
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(data[0].keys()):
                    hdr_cells[i].text = str(header)
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                
                # Add rows
                for row_data in data:
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row_data.values()):
                        row_cells[i].text = str(value)
            
            # Save file
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_filename = "".join(c for c in filename if c.isalnum() or c in (' ', '_', '-'))
            file_path = f"/tmp/{safe_filename}_{timestamp}.docx"
            doc.save(file_path)
            
            return {
                "success": True,
                "file_path": file_path,
                "filename": os.path.basename(file_path),
                "format": "docx"
            }
        
        except Exception as e:
            logger.error(f"Error exporting to Word: {str(e)}")
            return {"success": False, "error": str(e)}
    
    @frappe.whitelist()
    def export_to_pdf(self, title: str, data: List[Dict], filename: str) -> Dict:
        """Export data to PDF document"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
            from reportlab.lib import colors
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_filename = "".join(c for c in filename if c.isalnum() or c in (' ', '_', '-'))
            file_path = f"/tmp/{safe_filename}_{timestamp}.pdf"
            
            # Create PDF
            pdf = SimpleDocTemplate(file_path, pagesize=letter)
            elements = []
            styles = getSampleStyleSheet()
            
            # Add title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1f4788'),
                spaceAfter=30,
                alignment=1  # Center
            )
            elements.append(Paragraph(title, title_style))
            elements.append(Spacer(1, 0.3*inch))
            
            # Add metadata
            elements.append(Paragraph(
                f"<i>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</i>",
                styles['Normal']
            ))
            elements.append(Spacer(1, 0.2*inch))
            
            # Add table
            if data:
                table_data = [list(data[0].keys())] + [list(row.values()) for row in data]
                
                table = Table(table_data, colWidths=[2*inch, 2*inch, 1.5*inch, 1.5*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                elements.append(table)
            
            # Build PDF
            pdf.build(elements)
            
            return {
                "success": True,
                "file_path": file_path,
                "filename": os.path.basename(file_path),
                "format": "pdf"
            }
        
        except Exception as e:
            logger.error(f"Error exporting to PDF: {str(e)}")
            return {"success": False, "error": str(e)}
    
    @frappe.whitelist()
    def send_report_email(self, recipients: str, subject: str, message: str,
                         file_path: str, file_format: str = "pdf") -> Dict:
        """Send report via email with attachment"""
        try:
            from frappe.core.doctype.communication import make as make_email
            import mimetypes
            
            # Prepare attachment
            with open(file_path, 'rb') as f:
                file_content = f.read()
            
            filename = os.path.basename(file_path)
            
            # Send email
            email = frappe.get_doc({
                "doctype": "Email",
                "subject": subject,
                "recipients": recipients,
                "message": message,
            })
            
            # Add attachment
            email.add_attachment(
                fname=filename,
                fcontent=file_content,
                ftype=mimetypes.guess_type(filename)[0] or 'application/octet-stream'
            )
            
            email.send()
            
            return {
                "success": True,
                "message": f"Report sent to {recipients}",
                "email_id": email.name
            }
        
        except Exception as e:
            logger.error(f"Error sending email: {str(e)}")
            return {"success": False, "error": str(e)}
    
    @frappe.whitelist()
    def export_sales_report(self, days: int = 30, format: str = "xlsx") -> Dict:
        """Generate and export complete sales report"""
        try:
            date_from = frappe.utils.add_days(frappe.utils.today(), -days)
            
            # Fetch data
            sales = frappe.db.get_list(
                "Sales Order",
                filters=[
                    ["creation", ">=", date_from],
                    ["status", "!=", "Cancelled"]
                ],
                fields=["name", "customer", "grand_total", "status", "creation"],
                limit_page_length=1000
            )
            
            # Export based on format
            if format == "xlsx":
                return self.export_to_excel(sales, "Sales_Report", "Sales Orders")
            elif format == "pdf":
                return self.export_to_pdf("Sales Report", sales, "Sales_Report")
            elif format == "docx":
                return self.export_to_word("Sales Report", sales, "Sales_Report")
            else:
                return {"success": False, "error": "Invalid format"}
        
        except Exception as e:
            logger.error(f"Error exporting sales report: {str(e)}")
            return {"success": False, "error": str(e)}


# Public API endpoints
@frappe.whitelist()
def export_to_excel(data, filename, sheet_name="Report"):
    manager = ExportManager()
    data = json.loads(data) if isinstance(data, str) else data
    return manager.export_to_excel(data, filename, sheet_name)


@frappe.whitelist()
def export_to_pdf(title, data, filename):
    manager = ExportManager()
    data = json.loads(data) if isinstance(data, str) else data
    return manager.export_to_pdf(title, data, filename)


@frappe.whitelist()
def send_report_email(recipients, subject, message, file_path, file_format="pdf"):
    manager = ExportManager()
    return manager.send_report_email(recipients, subject, message, file_path, file_format)


import json
